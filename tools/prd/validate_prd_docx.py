from pathlib import Path
from docx import Document
import sys

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "artifacts/Web3 Remote Job Copilot 个人求职冲刺版 PRD.docx"

REQUIRED_TEXT = [
    "Web3 Remote Job Copilot 个人求职冲刺版 PRD",
    "Mia Web3 Remote Application Command Center",
    "V1 是一个单用户 Web App",
    "第一用户是 Mia",
    "30 天主冲刺",
    "60 天转化窗口",
    "每周 review 80-120 个岗位",
    "每周完成 20-30 个高质量投递",
    "每周发送 30-50 条定向 outreach",
    "30 天内获得 3-5 个有效回复",
    "60 天目标：拿到 1 个 remote offer",
    "Candidate Asset Layer",
    "Today Command Center",
    "Application Pack Builder",
    "Outreach Tracker",
    "Fit & Risk Score",
    "Weekly Review",
    "Growth Data Analyst",
    "Business Analyst",
    "Product / Operations Analyst",
    "Research & Due Diligence Analyst",
    "不自动登录 LinkedIn / Indeed",
    "不自动发送连接请求、DM 或申请",
    "不自动提交申请",
    "所有外部动作都必须人审确认，尤其是申请、DM、follow-up。",
    "少自动化投递，多提高命中率。",
    "允许用户手动粘贴职位 URL、联系人 URL、JD 文本。",
    "不自动抓取登录态页面。",
    "申请提交默认走人审确认。",
    "允许提高 Web3-adjacent remote 岗位权重，以增加上岸概率。",
    "Greenhouse",
    "Lever",
    "Remotive",
]

FORBIDDEN_TEXT = [
    "本产品的核心是自动海投",
    "支持自动登录 LinkedIn",
    "支持自动登录 LinkedIn / Indeed",
    "可自动登录 LinkedIn",
    "可自动登录 LinkedIn / Indeed",
    "支持自动发送 LinkedIn DM",
    "支持自动提交 Indeed Apply",
    "V1 支持多用户账号",
    "V1 支持订阅支付",
    "V1 支持团队后台",
    "V1 支持浏览器插件",
    "V1 支持 LinkedIn 自动化",
    "V1 支持 Indeed 自动化",
    "V1 支持自动提交申请",
]


def fail(message: str) -> None:
    print(f"FAIL: {message}")
    sys.exit(1)


def extract_text(document: Document) -> str:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def main() -> None:
    if not DOCX.exists():
        fail(f"missing docx: {DOCX}")

    document = Document(DOCX)
    text = extract_text(document)

    for phrase in REQUIRED_TEXT:
        if phrase not in text:
            fail(f"missing required text: {phrase}")

    for phrase in FORBIDDEN_TEXT:
        if phrase in text:
            fail(f"forbidden text present: {phrase}")

    if len(document.paragraphs) < 80:
        fail(f"expected at least 80 paragraphs, got {len(document.paragraphs)}")

    if len(document.tables) < 1:
        fail("expected at least one table")

    print("PASS: generated PRD docx validates")


if __name__ == "__main__":
    main()
