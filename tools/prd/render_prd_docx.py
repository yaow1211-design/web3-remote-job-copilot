from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs/prd/web3_remote_application_command_center_prd.md"
OUTPUT = ROOT / "artifacts/Web3 Remote Job Copilot 个人求职冲刺版 PRD.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(10)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.style = "Intense Quote"
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_paragraph_with_inline_bold(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)

    remaining = text
    while "**" in remaining:
        before, marker, after = remaining.partition("**")
        if before:
            paragraph.add_run(before)
        bold_text, marker2, tail = after.partition("**")
        if marker2:
            run = paragraph.add_run(bold_text)
            run.bold = True
            remaining = tail
        else:
            paragraph.add_run(marker + after)
            remaining = ""
    if remaining:
        paragraph.add_run(remaining)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(12)


def render() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)

    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_markdown_table(document, table_lines)
            table_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            add_code_block(document, code_lines)
            code_lines = []

    for raw_line in markdown:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            paragraph = document.add_heading(line[2:].strip(), level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
            continue

        if line[:3].strip(".").isdigit() and ". " in line[:5]:
            number, _, body = line.partition(". ")
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(body.strip())
            continue

        add_paragraph_with_inline_bold(document, line)

    flush_table()
    flush_code()
    document.save(OUTPUT)
    print(f"Rendered {OUTPUT}")


if __name__ == "__main__":
    render()
